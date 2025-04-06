import requests
from bs4 import BeautifulSoup
import re
from docx import Document

# 目標 URL
url = 'https://www.icrt.com.tw/news_lunchbox.php?mlevel1=7&mlevel2=96'

headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
}

# 發送 GET 請求
response = requests.get(url, headers=headers)
response.encoding = 'utf-8'  # 若必要，可設定編碼

# 解析 HTML
soup = BeautifulSoup(response.text, 'html.parser')

# 擷取日期，格式例：發布於 2025-02-12
date_div = soup.find('div', class_='news_lunchBox_date')
date_str = date_div.get_text(strip=True) if date_div else "Unknown Date"

# 準備儲存不同區塊內容的字典
sections = {}

# 找出所有 news_lunchBox_audio 區塊
audio_divs = soup.find_all('div', class_='news_lunchBox_audio')

for audio in audio_divs:
    title_div = audio.find('div', class_='styD_box')
    if not title_div:
        continue
    title = title_div.get_text(strip=True)
    
    txts_div = audio.find('div', class_='txts')
    if not txts_div:
        continue

    content_parts = []
    for item in txts_div.contents:
        if item.name == 'br':
            content_parts.append('\n\n')
        else:
            text = item.strip() if isinstance(item, str) else item.get_text(strip=True)
            content_parts.append(text)

    content = ''.join(content_parts)
    content = re.sub(r'\n+', '\n', content).strip()

    # 找到 "Quiz" 之前的內容
    cut_off_marker1 = "________________________________________\nQuiz"
    cut_off_marker2 = "-----------------------------------------------------------------------------------------------------------\nQuiz"
    if cut_off_marker1 in content:
        content = content.split(cut_off_marker1)[0].strip()
    
    if cut_off_marker2 in content:
        content = content.split(cut_off_marker2)[0].strip()

    sections[title] = content

# 建立 Word 文件
doc = Document()

# 加入日期作為 Level 1 標題
doc.add_heading(date_str, level=1)

# 加入不同區塊
if "News for Kids（國小）" in sections:
    doc.add_heading("News for Kids（國小）", level=2)
    doc.add_paragraph(sections["News for Kids（國小）"])
if "News Bites（國中）" in sections:
    doc.add_heading("News Bites（國中）", level=2)
    doc.add_paragraph(sections["News Bites（國中）"])

# 儲存為 docx 檔案
doc.save("ICRT_News.docx")

print("資料已整理並輸出到 ICRT_News.docx")
